#!/usr/bin/env python3
import json
import numpy as np
import pandas as pd
from docx import Document

totr = 153944
mortr = 30194
# import data dict json
f = open("data/data_dictionary.json")
data = json.load(f)
# remove cause of death
cod = data.pop("DEATH_CODE")
names_num = [(data[k]["name"], k) for k in data.keys()
             if "name" in data[k].keys() and
             data[k]["type"] == "numerical" and
             "summary" in data[k].keys()]
names_cat = [(data[k]["name"], k) for k in data.keys()
             if "name" in data[k].keys() and
             data[k]["type"] == "categorical" and
             "summary" in data[k].keys()]

toc = ['AGE', 'Sex', 'BMI', 'Race1', 'SMO_H', 'DB', 'ARRT', 'CHF', 'NYHA', 'EF', 'ICU', 'VENT', 'TP', 'IE', 'STAT', 'CCT', 'PERF']
print(len(data.keys()))
# print(names_cat)
# define variables

doc = Document()

# helper functions
def dict_to_str(d, n=None, den=1.0):
    out = ""
    for k in d:
        if n is None:
            out += "{} ({:.2f}%)\n".format(d[k], 100*d[k]/den)
        else:
            out += "- {}\n".format(k)
    return out

# add table for numerical valuse
# cols [name, mean_alive (std), mean_dead (std), pvalue]
tablet = doc.add_table(rows=1, cols=4)

hdr_cells = tablet.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'Survived (n=123750): Mean (SD) or Count (%)'
hdr_cells[2].text = 'Mortality (n=30194): Mean (SD) or Count (%)'
hdr_cells[3].text = 'p-value'

for k in toc:
    d = data[k]["summary"]
    p = "<0.001" if d["pvalue"] < 0.001 else "{:.2f}".format(d["pvalue"])
    if data[k]["type"] == "numerical":
        row_cells = tablet.add_row().cells
        row_cells[0].text = data[k]["name"]
        row_cells[1].text = "{:.2f} ({:.2f})".format(d["mean_alive"], d["std_alive"])
        row_cells[2].text = "{:.2f} ({:.2f})".format(d["mean_mort"], d["std_mort"])
        row_cells[3].text = p
    else:
        row_cells1 = tablet.add_row().cells
        row_cells2 = tablet.add_row().cells
        p = "<0.001" if d["pvalue"] < 0.001 else "{:.2f}".format(d["pvalue"])
        row_cells1[0].text = data[k]["name"]
        row_cells2[0].text = dict_to_str(d["counts"], n=data[k]["name"])
        row_cells2[1].text = dict_to_str(d["counts_alive"], den=totr-mortr)
        row_cells2[2].text = dict_to_str(d["counts_mort"], den=mortr)
        row_cells2[3].text = p
        row_cells1[0].merge(row_cells1[1].merge(row_cells1[2].merge(row_cells1[3])))


doc.add_page_break()

# add table for numerical valuse
# cols [name, mean_alive (std), mean_dead (std), pvalue]
table1 = doc.add_table(rows=1, cols=4)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'Survived Mean (SD)'
hdr_cells[2].text = 'Mortality Mean (SD)'
hdr_cells[3].text = 'p-value'

for rec in names_num:
    d = data[rec[1]]["summary"]
    row_cells = table1.add_row().cells
    p = "<0.001" if d["pvalue"] < 0.001 else "{:.2f}".format(d["pvalue"])
    row_cells[0].text = rec[0]
    row_cells[1].text = "{:.2f} ({:.2f})".format(d["mean_alive"], d["std_alive"])
    row_cells[2].text = "{:.2f} ({:.2f})".format(d["mean_mort"], d["std_alive"])
    row_cells[3].text = p

doc.add_page_break()

table2 = doc.add_table(rows=1, cols=4)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'Survived Counts'
hdr_cells[2].text = 'Mortality Counts'
hdr_cells[3].text = 'p-value'


for rec in names_cat:
    d = data[rec[1]]["summary"]
    row_cells1 = table2.add_row().cells
    row_cells2 = table2.add_row().cells
    p = "<0.001" if d["pvalue"] < 0.001 else "{:.2f}".format(d["pvalue"])
    row_cells1[0].text = rec[0]
    row_cells2[0].text = dict_to_str(d["counts"], n=rec[0])
    row_cells2[1].text = dict_to_str(d["counts_alive"])
    row_cells2[2].text = dict_to_str(d["counts_mort"])
    row_cells2[3].text = p

doc.save('tables.docx')

